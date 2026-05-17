"""DOCX export helpers."""

from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Inches

from webinar_transcriber.export.formatting import section_timecode
from webinar_transcriber.text_utils import split_paragraph_blocks

if TYPE_CHECKING:
    from collections.abc import Callable

    from docx.document import Document as DocxDocument

    from webinar_transcriber.models import ReportDocument, ReportSection

_FORMATTED_PARAGRAPH_RE = re.compile(
    r"^(?:"
    r"\*\*(?P<speaker>S\d+):\*\*\s*(?P<speaker_body>.*)"
    r"|[-*]\s+(?P<bullet>.+)"
    r"|\d+[.)]\s+(?P<number>.+)"
    r")$"
)


def write_docx_report(
    report: ReportDocument,
    output_path: Path,
    *,
    warning_callback: Callable[[str], None] | None = None,
) -> Path:
    """Write the report to DOCX.

    Returns:
        Path: The written DOCX artifact path.
    """
    document = Document()
    document.add_heading(report.title, level=0)

    if report.detected_language:
        document.add_paragraph(f"Language: {report.detected_language}")

    if report.summary:
        document.add_heading("Summary", level=1)
        for item in report.summary:
            document.add_paragraph(item, style="List Bullet")

    if report.action_items:
        document.add_heading("Action Items", level=1)
        for item in report.action_items:
            document.add_paragraph(item, style="List Bullet")

    document.add_heading("Sections", level=1)
    for section in report.sections:
        _add_section(
            document,
            section,
            image_base_dir=output_path.parent,
            warning_callback=warning_callback,
        )

    document.save(str(output_path))
    return output_path


def _add_section(
    document: DocxDocument,
    section: ReportSection,
    *,
    image_base_dir: Path,
    warning_callback: Callable[[str], None] | None = None,
) -> None:
    title = section.title
    timecode = section_timecode(section.start_sec, section.end_sec)
    document.add_heading(f"{title} ({timecode})", level=2)
    _add_section_image(
        document,
        section.image_path,
        image_base_dir=image_base_dir,
        warning_callback=warning_callback,
    )
    _add_section_tldr(document, section.tldr)
    paragraphs = split_paragraph_blocks(section.transcript_text) or [section.transcript_text]
    for paragraph_text in paragraphs:
        if not _add_formatted_paragraph(document, paragraph_text):
            document.add_paragraph(paragraph_text)


def _add_section_image(
    document: DocxDocument,
    image_path: str | None,
    *,
    image_base_dir: Path,
    warning_callback: Callable[[str], None] | None = None,
) -> None:
    if not image_path:
        return

    resolved_path = Path(image_path)
    if not resolved_path.is_absolute():
        resolved_path = image_base_dir / resolved_path
    if not resolved_path.exists():
        if warning_callback is not None:
            warning_callback(f"Section image does not exist: {resolved_path}")
        return
    document.add_picture(str(resolved_path), width=Inches(6))


def _add_section_tldr(document: DocxDocument, tldr: str | None) -> None:
    if not tldr or not (text := tldr.strip()):
        return

    tldr_label = document.add_paragraph()
    tldr_label.add_run("TL;DR / Cheat Sheet").bold = True
    _add_text_blocks(document, text)
    transcript_label = document.add_paragraph()
    transcript_label.add_run("Transcript").bold = True


def _add_text_blocks(document: DocxDocument, text: str) -> None:
    paragraphs = split_paragraph_blocks(text) or [text]
    for paragraph_text in paragraphs:
        lines = [line.strip() for line in paragraph_text.splitlines() if line.strip()]
        for line in lines:
            if _add_formatted_paragraph(document, line):
                continue
            document.add_paragraph(line)


def _add_formatted_paragraph(document: DocxDocument, text: str) -> bool:
    match = _FORMATTED_PARAGRAPH_RE.match(text)
    if match is None:
        return False

    if speaker := match.group("speaker"):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{speaker}:").bold = True
        if body := (match.group("speaker_body") or "").strip():
            paragraph.add_run(f" {body}")
        return True

    if body := match.group("bullet"):
        document.add_paragraph(body, style="List Bullet")
        return True

    document.add_paragraph(match.group("number") or "", style="List Number")
    return True
