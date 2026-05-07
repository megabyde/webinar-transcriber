"""Tests for report exporters."""

import json
from pathlib import Path

from docx import Document
from PIL import Image

from webinar_transcriber.export.docx_report import write_docx_report
from webinar_transcriber.export.formatting import format_timecode
from webinar_transcriber.export.json_report import write_json_report
from webinar_transcriber.export.markdown import write_markdown_report
from webinar_transcriber.models import MediaType, ReportDocument, ReportSection


def _style_name(paragraph) -> str | None:
    style = paragraph.style
    return None if style is None else style.name


class TestFormatting:
    def test_format_timecode_formats_short_and_long_durations(self) -> None:
        assert format_timecode(65.9) == "01:05"
        assert format_timecode(3661.2) == "01:01:01"


class TestDocxReport:
    def test_splits_blank_line_paragraphs(self, tmp_path: Path) -> None:
        report = ReportDocument(
            title="Demo",
            source_file="demo.wav",
            media_type=MediaType.AUDIO,
            sections=[
                ReportSection(
                    id="section-1",
                    title="Section 1",
                    start_sec=0.0,
                    end_sec=5.0,
                    tldr="Краткое резюме раздела.",
                    transcript_text="Первый абзац.\n\nВторой абзац.",  # noqa: RUF001
                )
            ],
        )

        output_path = tmp_path / "report.docx"
        write_docx_report(report, output_path)

        document = Document(str(output_path))
        paragraph_texts = [paragraph.text for paragraph in document.paragraphs]

        assert "Section 1 (00:00\N{EN DASH}00:05)" in paragraph_texts
        assert "TL;DR / Cheat Sheet" in paragraph_texts
        assert "Transcript" in paragraph_texts
        assert "Краткое резюме раздела." in paragraph_texts
        assert "Первый абзац." in paragraph_texts
        assert "Второй абзац." in paragraph_texts

    def test_formats_cheat_sheet_lists(self, tmp_path: Path) -> None:
        report = ReportDocument(
            title="Demo",
            source_file="demo.wav",
            media_type=MediaType.AUDIO,
            sections=[
                ReportSection(
                    id="section-1",
                    title="Section 1",
                    start_sec=0.0,
                    end_sec=5.0,
                    tldr="- First point.\n\n1. Second point.\n2) Third point.",
                    transcript_text="Transcript body.",
                )
            ],
        )

        output_path = tmp_path / "report.docx"
        write_docx_report(report, output_path)

        document = Document(str(output_path))
        paragraph_data = [
            (paragraph.text, _style_name(paragraph)) for paragraph in document.paragraphs
        ]

        assert ("First point.", "List Bullet") in paragraph_data
        assert ("Second point.", "List Number") in paragraph_data
        assert ("Third point.", "List Number") in paragraph_data

    def test_formats_canonical_tldr_lines(self, tmp_path: Path) -> None:
        report = ReportDocument(
            title="Demo",
            source_file="demo.wav",
            media_type=MediaType.AUDIO,
            sections=[
                ReportSection(
                    id="section-1",
                    title="Section 1",
                    start_sec=0.0,
                    end_sec=5.0,
                    tldr="\n\nHeading line\n- Bullet point.\n1. Numbered point\n\n\n",
                    transcript_text="Transcript body.",
                )
            ],
        )

        output_path = tmp_path / "report.docx"
        write_docx_report(report, output_path)

        document = Document(str(output_path))
        paragraph_data = [
            (paragraph.text, _style_name(paragraph)) for paragraph in document.paragraphs
        ]

        assert ("Heading line", "Normal") in paragraph_data
        assert ("Bullet point.", "List Bullet") in paragraph_data
        assert ("Numbered point", "List Number") in paragraph_data
        assert ("Transcript", "Normal") in paragraph_data

    def test_skips_missing_section_image(self, tmp_path: Path) -> None:
        report = ReportDocument(
            title="Demo",
            source_file="demo.wav",
            media_type=MediaType.AUDIO,
            sections=[
                ReportSection(
                    id="section-1",
                    title="Section 1",
                    start_sec=0.0,
                    end_sec=5.0,
                    transcript_text="Paragraph",
                    image_path=str(tmp_path / "missing.png"),
                )
            ],
        )

        output_path = tmp_path / "report.docx"
        warnings: list[str] = []

        write_docx_report(report, output_path, warning_callback=warnings.append)

        assert output_path.exists()
        assert warnings == [f"Section image does not exist: {tmp_path / 'missing.png'}"]

    def test_embeds_section_image(self, tmp_path: Path) -> None:
        image_path = tmp_path / "frame.png"
        Image.new("RGB", (8, 8), color="white").save(image_path)
        report = ReportDocument(
            title="Demo",
            source_file="demo.wav",
            media_type=MediaType.VIDEO,
            sections=[
                ReportSection(
                    id="section-1",
                    title="Section 1",
                    start_sec=0.0,
                    end_sec=5.0,
                    transcript_text="Paragraph",
                    image_path=str(image_path),
                )
            ],
        )

        output_path = tmp_path / "report.docx"
        write_docx_report(report, output_path)

        document = Document(str(output_path))
        assert len(document.inline_shapes) == 1

    def test_embeds_relative_section_image_from_report_directory(self, tmp_path: Path) -> None:
        frames_dir = tmp_path / "frames"
        frames_dir.mkdir()
        Image.new("RGB", (8, 8), color="white").save(frames_dir / "scene-1.png")
        report = ReportDocument(
            title="Demo",
            source_file="demo.mp4",
            media_type=MediaType.VIDEO,
            sections=[
                ReportSection(
                    id="section-1",
                    title="Section 1",
                    start_sec=0.0,
                    end_sec=5.0,
                    transcript_text="Paragraph",
                    image_path="frames/scene-1.png",
                )
            ],
        )

        output_path = tmp_path / "report.docx"
        write_docx_report(report, output_path)

        document = Document(str(output_path))
        assert len(document.inline_shapes) == 1

    def test_blank_tldr_does_not_add_extra_body_paragraphs(self, tmp_path: Path) -> None:
        report = ReportDocument(
            title="Demo",
            source_file="demo.wav",
            media_type=MediaType.AUDIO,
            sections=[
                ReportSection(
                    id="section-1",
                    title="Section 1",
                    start_sec=0.0,
                    end_sec=5.0,
                    tldr="  \n  ",
                    transcript_text="Transcript body.",
                )
            ],
        )

        output_path = tmp_path / "report.docx"
        write_docx_report(report, output_path)

        document = Document(str(output_path))
        assert [paragraph.text for paragraph in document.paragraphs] == [
            "Demo",
            "Sections",
            "Section 1 (00:00\N{EN DASH}00:05)",
            "TL;DR / Cheat Sheet",
            "Transcript",
            "Transcript body.",
        ]


class TestMarkdownReport:
    def test_omits_blank_image_line_for_imageless_sections(self, tmp_path: Path) -> None:
        report = ReportDocument(
            title="Demo",
            source_file="demo.wav",
            media_type=MediaType.AUDIO,
            detected_language="en",
            summary=["Summary point."],
            action_items=["Follow up."],
            sections=[
                ReportSection(
                    id="section-1",
                    title="Section 1",
                    start_sec=0.0,
                    end_sec=5.0,
                    tldr="Section recap.",
                    transcript_text="Paragraph one.",
                ),
                ReportSection(
                    id="section-2",
                    title="Section 2",
                    start_sec=5.0,
                    end_sec=10.0,
                    transcript_text="Paragraph two.",
                    image_path="frames/scene-2.png",
                ),
            ],
        )

        output_path = tmp_path / "report.md"
        write_markdown_report(report, output_path)

        expected_markdown = (
            "# Demo\n\n"
            "- Language: `en`\n\n"
            "## Summary\n\n"
            "- Summary point.\n\n"
            "## Action Items\n\n"
            "- Follow up.\n\n"
            "## Sections\n\n"
            "### Section 1 (00:00\N{EN DASH}00:05)\n\n"
            "**TL;DR / Cheat Sheet**\n\n"
            "Section recap.\n\n"
            "**Transcript**\n\n"
            "Paragraph one.\n\n"
            "### Section 2 (00:05\N{EN DASH}00:10)\n\n"
            "![Section 2](frames/scene-2.png)\n\n"
            "Paragraph two.\n"
        )

        assert output_path.read_text(encoding="utf-8") == expected_markdown

    def test_uses_hour_format_for_long_sections(self, tmp_path: Path) -> None:
        report = ReportDocument(
            title="Demo",
            source_file="demo.wav",
            media_type=MediaType.AUDIO,
            sections=[
                ReportSection(
                    id="section-1",
                    title="Section 1",
                    start_sec=3661.2,
                    end_sec=7325.9,
                    transcript_text="Paragraph one.",
                )
            ],
        )

        output_path = tmp_path / "report.md"
        write_markdown_report(report, output_path)

        markdown = output_path.read_text(encoding="utf-8")

        assert "### Section 1 (01:01:01\N{EN DASH}02:02:05)" in markdown


class TestJsonReport:
    def test_round_trips_report_document(self, tmp_path: Path) -> None:
        report = ReportDocument(
            title="Demo",
            source_file="demo.wav",
            media_type=MediaType.AUDIO,
            summary=["Summary point."],
            action_items=["Follow up."],
            sections=[
                ReportSection(
                    id="section-1",
                    title="Section 1",
                    start_sec=0.0,
                    end_sec=5.0,
                    transcript_text="Paragraph one.",
                )
            ],
        )

        output_path = tmp_path / "report.json"
        write_json_report(report, output_path)

        payload = json.loads(output_path.read_text(encoding="utf-8"))

        assert payload["title"] == "Demo"
        assert payload["media_type"] == "audio"
        assert payload["summary"] == ["Summary point."]
        assert payload["action_items"] == ["Follow up."]
        assert payload["sections"] == [
            {
                "id": "section-1",
                "title": "Section 1",
                "start_sec": 0.0,
                "end_sec": 5.0,
                "transcript_text": "Paragraph one.",
                "tldr": None,
                "frame_id": None,
                "image_path": None,
            }
        ]
