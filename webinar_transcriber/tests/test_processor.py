"""Tests for the end-to-end processor flow."""

import json
from collections.abc import Callable
from pathlib import Path
from types import SimpleNamespace
from typing import cast
from unittest.mock import patch

import numpy as np
import pytest
from docx import Document

from webinar_transcriber.asr import WhisperCppTranscriber
from webinar_transcriber.llm import (
    LLMConfigurationError,
    LLMProcessingError,
    LLMProcessor,
    LLMReportMetadataResult,
    LLMReportPolishPlan,
    LLMSectionPolishResult,
)
from webinar_transcriber.models import (
    DecodedWindow,
    MediaType,
    ReportDocument,
    ReportSection,
    SpeechRegion,
    TranscriptSegment,
    VideoAsset,
)
from webinar_transcriber.processor import (
    ProcessArtifacts,
    process_input,
)
from webinar_transcriber.processor.__init__ import (
    _AsrPipelineState,
    _RunContext,
    _write_run_diagnostics,
)
from webinar_transcriber.processor.llm import LLMRuntimeState, resolve_llm_processor
from webinar_transcriber.processor.support import (
    asr_model_label,
    asr_runtime_detail,
    hf_cache_repo_label,
    title_update_detail,
    token_usage_detail,
    window_transcription_stage_detail,
)
from webinar_transcriber.reporter import NullStageReporter

FIXTURE_DIR = Path(__file__).parent / "fixtures"
EXPECTED_LLM_WARNING = (
    "Section polish response returned an empty transcript text for section-1; kept original text."
)
EXPECTED_LLM_USAGE = {"input_tokens": 20, "output_tokens": 6, "total_tokens": 26}
EXPECTED_LLM_SECTION_TEXT = (
    "Agenda review, and project status update.\n\nNext step: please send the draft by Friday."
)


def install_basic_windowing(
    monkeypatch: pytest.MonkeyPatch,
    *,
    region_end_sec: float = 6.0,
    sample_rate: int = 16_000,
    load_audio: Callable[[Path], tuple[np.ndarray, int]] | None = None,
) -> None:
    """Patch processor seams so tests exercise the windowed flow deterministically."""

    if load_audio is None:

        def load_audio(_path: Path) -> tuple[np.ndarray, int]:
            return np.zeros(sample_rate, dtype=np.float32), sample_rate

    monkeypatch.setattr("webinar_transcriber.processor.asr.load_normalized_audio", load_audio)
    monkeypatch.setattr(
        "webinar_transcriber.processor.asr.detect_speech_regions",
        lambda *_args, **_kwargs: ([SpeechRegion(start_sec=0.0, end_sec=region_end_sec)], []),
    )


class RecordingReporter(NullStageReporter):
    """Collect stage updates for assertions."""

    def __init__(self) -> None:
        self.events: list[tuple[str, str, str]] = []
        self.warnings: list[str] = []
        self.progress_events: list[tuple[str, str, float, str | None]] = []

    def begin_run(self, input_path: Path) -> None:
        self.events.append(("begin", input_path.name, ""))

    def stage_started(self, stage_key: str, label: str) -> None:
        self.events.append(("start", stage_key, label))

    def progress_started(
        self,
        stage_key: str,
        label: str,
        *,
        total: float,
        count_label: str | None = None,
        count_multiplier: float = 1.0,
        rate_label: str | None = None,
        rate_multiplier: float = 1.0,
        detail: str | None = None,
    ) -> None:
        self.events.append(("start", stage_key, label))
        self.progress_events.append(("start", stage_key, total, detail))

    def progress_advanced(
        self, stage_key: str, *, advance: float = 1.0, detail: str | None = None
    ) -> None:
        self.progress_events.append(("advance", stage_key, advance, detail))

    def stage_finished(self, stage_key: str, label: str, *, detail: str | None = None) -> None:
        self.events.append(("finish", stage_key, detail or ""))

    def warn(self, message: str) -> None:
        self.warnings.append(message)

    def complete_run(self, artifacts: ProcessArtifacts) -> None:
        self.events.append((
            "complete",
            artifacts.layout.run_dir.name,
            artifacts.report.source_file,
        ))

    def has_event(self, event_type: str, stage_key: str, detail: str) -> bool:
        return (event_type, stage_key, detail) in self.events

    def has_event_detail(
        self, event_type: str, stage_key: str, predicate: Callable[[str], bool]
    ) -> bool:
        return any(
            event[0] == event_type and event[1] == stage_key and predicate(event[2])
            for event in self.events
        )

    def has_progress_event(
        self, event_type: str, stage_key: str, value: float, detail: str | None
    ) -> bool:
        return (event_type, stage_key, value, detail) in self.progress_events

    def has_progress_event_detail(
        self, event_type: str, stage_key: str, predicate: Callable[[str | None], bool]
    ) -> bool:
        return any(
            event[0] == event_type and event[1] == stage_key and predicate(event[3])
            for event in self.progress_events
        )

    def progress_stage_events(
        self, event_type: str, stage_key: str
    ) -> list[tuple[str, str, float, str | None]]:
        return [
            event
            for event in self.progress_events
            if event[0] == event_type and event[1] == stage_key
        ]


class TestProcessInput:
    class FakeTranscriber(WhisperCppTranscriber):
        """Stable whisper-style test double for deterministic transcripts."""

        def __init__(
            self, *, detected_language: str = "en", segments: list[TranscriptSegment] | None = None
        ) -> None:
            super().__init__(model_name="test-model")
            self._detected_language = detected_language
            self._segments = segments or [
                TranscriptSegment(
                    id="segment-1",
                    text="Agenda review and project status update.",
                    start_sec=0.0,
                    end_sec=3.0,
                ),
                TranscriptSegment(
                    id="segment-2",
                    text="Next step please send the draft by Friday.",
                    start_sec=3.0,
                    end_sec=6.0,
                ),
            ]

        @property
        def device_name(self) -> str:
            return "cpu"

        @property
        def system_info(self) -> str:
            return "CPU = 1"

        def prepare_model(self) -> None:
            """No-op test hook for the ASR preparation stage."""

        def transcribe_inference_windows(
            self, audio_samples, windows, *, progress_callback=None
        ) -> list[DecodedWindow]:
            del audio_samples
            assert progress_callback is not None
            for window in windows:
                progress_callback(window.end_sec, len(self._segments))
            return [
                DecodedWindow(
                    window=w,
                    text=" ".join(segment.text for segment in self._segments),
                    language=self._detected_language,
                    segments=list(self._segments),
                )
                for w in windows
            ]

    def test_writes_reports_and_metadata(self, tmp_path, monkeypatch) -> None:
        install_basic_windowing(monkeypatch)
        reporter = RecordingReporter()
        artifacts = process_input(
            FIXTURE_DIR / "sample-audio.mp3",
            output_dir=tmp_path / "run",
            transcriber=self.FakeTranscriber(),
            reporter=reporter,
        )

        assert artifacts.layout.metadata_path.exists()
        assert artifacts.layout.transcript_path.exists()
        assert artifacts.layout.subtitle_vtt_path.exists()
        assert artifacts.layout.markdown_report_path.exists()
        assert artifacts.layout.docx_report_path.exists()
        assert artifacts.layout.json_report_path.exists()
        assert not (artifacts.layout.run_dir / "audio.wav").exists()
        assert artifacts.report.detected_language == "en"
        assert artifacts.report.action_items == ["Next step please send the draft by Friday."]
        assert len(artifacts.report.sections) == 1
        assert (
            "Next step please send the draft by Friday."
            in artifacts.report.sections[0].transcript_text
        )
        assert artifacts.diagnostics.asr_backend == "whisper.cpp"
        assert artifacts.diagnostics.asr_model == "test-model"
        diagnostics_payload = json.loads(
            artifacts.layout.diagnostics_path.read_text(encoding="utf-8")
        )

        markdown = artifacts.layout.markdown_report_path.read_text(encoding="utf-8")
        vtt = artifacts.layout.subtitle_vtt_path.read_text(encoding="utf-8")
        assert "# Sample Audio" in markdown
        assert "Agenda review and project status update." in markdown
        assert "00:00:00.000 --> 00:00:03.000" in vtt
        assert "WEBVTT" in vtt
        assert reporter.has_event("start", "prepare_transcription_audio", "Preparing audio")
        assert reporter.has_event("finish", "prepare_transcription_audio", "sample-audio.wav")
        assert reporter.has_event("start", "probe_media", "Probing media")
        assert reporter.has_event("start", "prepare_asr", "Preparing ASR model")
        assert reporter.has_event("finish", "prepare_asr", "test-model | cpu")
        assert any(event[0] == "complete" for event in reporter.events)
        assert reporter.has_event_detail(
            "finish", "transcribe", lambda detail: "window" in detail and "RTF" in detail
        )
        transcribe_start_events = reporter.progress_stage_events("start", "transcribe")
        assert transcribe_start_events == [
            ("start", "transcribe", artifacts.media_asset.duration_sec, "0 segments")
        ]
        vad_start_events = reporter.progress_stage_events("start", "vad")
        assert vad_start_events == [("start", "vad", artifacts.media_asset.duration_sec, None)]
        assert reporter.has_progress_event_detail(
            "advance", "transcribe", lambda detail: detail == "2 segments"
        )
        structure_start_events = reporter.progress_stage_events("start", "structure")
        assert structure_start_events == [
            (
                "start",
                "structure",
                float(artifacts.diagnostics.item_counts["normalized_transcript_segments"]),
                "0 sections",
            )
        ]
        assert reporter.has_progress_event_detail(
            "advance", "structure", lambda detail: detail == "1 section"
        )
        assert reporter.has_progress_event_detail(
            "advance", "vad", lambda detail: detail == "1 region"
        )
        assert diagnostics_payload["asr_backend"] == "whisper.cpp"
        assert diagnostics_payload["asr_model"] == "test-model"

        document = Document(str(artifacts.layout.docx_report_path))
        assert "Sample Audio" in "\n".join(paragraph.text for paragraph in document.paragraphs)

    def test_write_run_diagnostics_returns_none_without_layout(self) -> None:
        ctx = _RunContext(
            reporter=NullStageReporter(),
            asr_pipeline=_AsrPipelineState(vad_enabled=True, threads=1),
        )

        diagnostics = _write_run_diagnostics(
            ctx,
            status="failed",
            failed_stage="prepare_run_dir",
            error="boom",
            asr_model="test-model",
            llm_enabled=False,
        )

        assert diagnostics is None

    def test_keeps_normalized_audio_artifact(self, tmp_path, monkeypatch) -> None:
        install_basic_windowing(monkeypatch)

        artifacts = process_input(
            FIXTURE_DIR / "sample-audio.mp3",
            output_dir=tmp_path / "run",
            transcriber=self.FakeTranscriber(),
            keep_audio=True,
            kept_audio_format="wav",
        )

        kept_audio_path = artifacts.layout.transcription_audio_path()

        assert kept_audio_path.exists()
        assert kept_audio_path.read_bytes()[:4] == b"RIFF"

    def test_does_not_close_caller_supplied_transcriber(self, tmp_path, monkeypatch) -> None:
        install_basic_windowing(monkeypatch)

        class CloseTrackingTranscriber(self.FakeTranscriber):
            def __init__(self) -> None:
                super().__init__()
                self.close_calls = 0

            def close(self) -> None:
                self.close_calls += 1

        transcriber = CloseTrackingTranscriber()

        process_input(
            FIXTURE_DIR / "sample-audio.mp3", output_dir=tmp_path / "run", transcriber=transcriber
        )

        assert transcriber.close_calls == 0

    def test_closes_owned_transcriber_on_failure(self, tmp_path, monkeypatch) -> None:
        install_basic_windowing(monkeypatch)

        class CloseTrackingTranscriber(self.FakeTranscriber):
            def __init__(self) -> None:
                super().__init__()
                self.close_calls = 0

            def close(self) -> None:
                self.close_calls += 1

        transcriber = CloseTrackingTranscriber()
        monkeypatch.setattr(
            "webinar_transcriber.asr.WhisperCppTranscriber", lambda *args, **kwargs: transcriber
        )
        monkeypatch.setattr(
            "webinar_transcriber.processor.asr.run_asr_pipeline",
            lambda **_kwargs: (_ for _ in ()).throw(RuntimeError("asr failed")),
        )

        with pytest.raises(RuntimeError, match="asr failed"):
            process_input(FIXTURE_DIR / "sample-audio.mp3", output_dir=tmp_path / "run")

        assert transcriber.close_calls == 1

    def test_writes_video_scene_artifacts(self, tmp_path, monkeypatch) -> None:
        reporter = RecordingReporter()
        transcription_audio_path: Path | None = None

        def _load_audio(audio_path: Path) -> tuple[np.ndarray, int]:
            nonlocal transcription_audio_path
            transcription_audio_path = audio_path
            assert audio_path.exists()
            return np.zeros(16_000, dtype=np.float32), 16_000

        install_basic_windowing(monkeypatch, region_end_sec=1.8, load_audio=_load_audio)
        artifacts = process_input(
            FIXTURE_DIR / "sample-video.mp4",
            output_dir=tmp_path / "video-run",
            transcriber=self.FakeTranscriber(
                segments=[
                    TranscriptSegment(
                        id="segment-1",
                        text="Agenda overview and open questions.",
                        start_sec=0.0,
                        end_sec=0.9,
                    ),
                    TranscriptSegment(
                        id="segment-2",
                        text="Timeline planning and next step review.",
                        start_sec=0.9,
                        end_sec=1.8,
                    ),
                ]
            ),
            reporter=reporter,
        )

        assert isinstance(artifacts.media_asset, VideoAsset)
        assert not (artifacts.layout.run_dir / "audio.wav").exists()
        assert transcription_audio_path is not None
        assert not transcription_audio_path.exists()
        assert artifacts.layout.scenes_path.exists()
        assert artifacts.layout.frames_dir.exists()
        assert any(artifacts.layout.frames_dir.iterdir())
        assert artifacts.report.sections
        assert artifacts.report.sections[0].image_path
        assert reporter.has_progress_event("start", "detect_scenes", 2, "0 scenes")
        assert reporter.has_progress_event_detail(
            "advance", "detect_scenes", lambda detail: detail == "1 scene" or detail == "2 scenes"
        )
        assert all(
            event[2] == 1.0
            for event in reporter.progress_events
            if event[0] == "advance" and event[1] == "detect_scenes"
        )
        assert reporter.has_progress_event(
            "start", "extract_frames", artifacts.diagnostics.item_counts["scenes"], None
        )

    def test_runs_windowed_whispercpp_pipeline(self, tmp_path, monkeypatch) -> None:
        reporter = RecordingReporter()

        class WindowedTranscriber(WhisperCppTranscriber):
            @property
            def system_info(self) -> str:
                return "METAL = 1"

            def prepare_model(self) -> None:
                return None

            def transcribe_inference_windows(
                self, audio_samples, windows, *, progress_callback=None
            ) -> list[DecodedWindow]:
                del audio_samples
                assert progress_callback is not None
                for window in windows:
                    progress_callback(window.end_sec, 2)
                return [
                    DecodedWindow(
                        window=windows[0],
                        text=(
                            "Agenda review and project status update. "
                            "Next step please send the draft by Friday."
                        ),
                        language="en",
                        segments=[
                            TranscriptSegment(
                                id="segment-1",
                                text="Agenda review and project status update.",
                                start_sec=0.0,
                                end_sec=3.0,
                            ),
                            TranscriptSegment(
                                id="segment-2",
                                text="Next step please send the draft by Friday.",
                                start_sec=3.0,
                                end_sec=6.0,
                            ),
                        ],
                    )
                ]

        monkeypatch.setattr(
            "webinar_transcriber.processor.asr.load_normalized_audio",
            lambda _path: (np.zeros(16_000, dtype=np.float32), 16_000),
        )
        monkeypatch.setattr(
            "webinar_transcriber.processor.asr.detect_speech_regions",
            lambda *_args, **_kwargs: ([SpeechRegion(start_sec=0.0, end_sec=6.0)], []),
        )
        artifacts = process_input(
            FIXTURE_DIR / "sample-audio.mp3",
            output_dir=tmp_path / "windowed-run",
            transcriber=WindowedTranscriber(model_name="test-model"),
            reporter=reporter,
        )

        assert artifacts.diagnostics.asr_pipeline is not None
        assert artifacts.diagnostics.asr_pipeline.window_count == 1
        assert artifacts.diagnostics.asr_pipeline.vad_region_count == 1
        assert artifacts.diagnostics.asr_pipeline.system_info == "METAL = 1"
        assert reporter.has_event("start", "vad", "Detecting speech regions")
        assert reporter.has_event("start", "prepare_speech_regions", "Preparing speech regions")
        assert reporter.has_event("start", "reconcile", "Reconciling transcript windows")
        assert artifacts.layout.speech_regions_path.exists()
        assert artifacts.layout.expanded_regions_path.exists()
        assert artifacts.layout.decoded_windows_path.exists()

        speech_regions_payload = json.loads(
            artifacts.layout.speech_regions_path.read_text(encoding="utf-8")
        )
        decoded_windows_payload = json.loads(
            artifacts.layout.decoded_windows_path.read_text(encoding="utf-8")
        )
        diagnostics_payload = json.loads(
            artifacts.layout.diagnostics_path.read_text(encoding="utf-8")
        )

        assert speech_regions_payload["speech_regions"][0] == {"start_sec": 0.0, "end_sec": 6.0}
        assert decoded_windows_payload["decoded_windows"][0]["window"]["window_id"] == "window-1"
        assert decoded_windows_payload["decoded_windows"][0]["language"] == "en"
        assert "input_prompt" in decoded_windows_payload["decoded_windows"][0]
        assert diagnostics_payload["item_counts"]["windows"] == 1

    def test_normalizes_transcript_before_report_generation(self, tmp_path, monkeypatch) -> None:
        install_basic_windowing(monkeypatch)
        reporter = RecordingReporter()

        artifacts = process_input(
            FIXTURE_DIR / "sample-audio.mp3",
            output_dir=tmp_path / "normalized-run",
            transcriber=self.FakeTranscriber(
                detected_language="ru",
                segments=[
                    TranscriptSegment(id="segment-1", text="", start_sec=0.0, end_sec=0.1),
                    TranscriptSegment(id="segment-2", text="Привет", start_sec=0.1, end_sec=0.8),
                    TranscriptSegment(id="segment-3", text="всем.", start_sec=0.8, end_sec=1.4),
                    TranscriptSegment(
                        id="segment-4",
                        text="Новая тема начинается позже.",
                        start_sec=3.0,
                        end_sec=6.0,
                    ),
                ],
            ),
            reporter=reporter,
        )

        assert artifacts.layout.transcript_path.exists()
        raw_payload = json.loads(artifacts.layout.transcript_path.read_text(encoding="utf-8"))
        segment_texts = [segment["text"] for segment in raw_payload["segments"]]

        assert segment_texts == ["Привет", "всем.", "Новая тема начинается позже."]
        assert artifacts.diagnostics.item_counts["normalized_transcript_segments"] == 2
        assert artifacts.report.sections[0].transcript_text.startswith("Привет всем.")

    def test_persists_intermediate_artifacts_on_failure(self, tmp_path) -> None:
        reporter = RecordingReporter()

        class WindowedFailureTranscriber(WhisperCppTranscriber):
            def prepare_model(self) -> None:
                return None

            def transcribe_inference_windows(
                self, audio_samples, windows, *, progress_callback=None
            ) -> list[DecodedWindow]:
                del audio_samples, progress_callback
                return [
                    DecodedWindow(
                        window=windows[0],
                        text="Agenda review and project status update.",
                        language="en",
                        segments=[
                            TranscriptSegment(
                                id="segment-1",
                                text="Agenda review and project status update.",
                                start_sec=0.0,
                                end_sec=3.0,
                            )
                        ],
                    )
                ]

        output_dir = tmp_path / "failed-run"
        with (
            patch("webinar_transcriber.structure.build_report", side_effect=RuntimeError("boom")),
            patch(
                "webinar_transcriber.processor.asr.load_normalized_audio",
                return_value=(np.zeros(16_000, dtype=np.float32), 16_000),
            ),
            patch(
                "webinar_transcriber.processor.asr.detect_speech_regions",
                return_value=([SpeechRegion(start_sec=0.0, end_sec=3.0)], []),
            ),
            pytest.raises(RuntimeError, match="boom"),
        ):
            process_input(
                FIXTURE_DIR / "sample-audio.mp3",
                output_dir=output_dir,
                transcriber=WindowedFailureTranscriber(model_name="test-model"),
                reporter=reporter,
            )

        assert (output_dir / "metadata.json").exists()
        assert (output_dir / "transcript.json").exists()
        assert (output_dir / "asr" / "speech_regions.json").exists()
        assert (output_dir / "asr" / "expanded_regions.json").exists()
        assert (output_dir / "asr" / "decoded_windows.json").exists()
        diagnostics_payload = json.loads(
            (output_dir / "diagnostics.json").read_text(encoding="utf-8")
        )
        assert diagnostics_payload["status"] == "failed"
        assert diagnostics_payload["failed_stage"] == "structure"
        assert diagnostics_payload["error"] == "boom"
        assert not (output_dir / "report.json").exists()

    def test_forwards_vad_warnings_to_reporter(self, tmp_path, monkeypatch) -> None:
        reporter = RecordingReporter()
        monkeypatch.setattr(
            "webinar_transcriber.processor.asr.load_normalized_audio",
            lambda _path: (np.zeros(16_000, dtype=np.float32), 16_000),
        )
        monkeypatch.setattr(
            "webinar_transcriber.processor.asr.detect_speech_regions",
            lambda *_args, **_kwargs: (
                [SpeechRegion(start_sec=0.0, end_sec=6.0)],
                ["Silero warning"],
            ),
        )

        process_input(
            FIXTURE_DIR / "sample-audio.mp3",
            output_dir=tmp_path / "warning-run",
            transcriber=self.FakeTranscriber(),
            reporter=reporter,
        )

        assert reporter.warnings == ["Silero warning"]


class TestProcessorHelpers:
    def test_window_transcription_stage_detail_reports_rtf(self) -> None:
        detail = window_transcription_stage_detail(
            window_count=1, total_duration_sec=12.5, elapsed_sec=2.5
        )

        assert detail == "1 window | RTF 0.20"

    def test_asr_runtime_detail_shortens_hf_cache_model_path(self) -> None:
        transcriber = TestProcessInput.FakeTranscriber()
        transcriber._model_path = Path(
            "/tmp/huggingface/hub/"
            "models--ggerganov--whisper.cpp/snapshots/123456/"
            "ggml-large-v3-turbo.bin"
        )

        assert asr_runtime_detail(transcriber) == (
            "ggerganov/whisper.cpp/ggml-large-v3-turbo.bin (HF cache) | cpu"
        )

    def test_title_update_detail_reports_partial_title_updates(self) -> None:
        assert title_update_detail(title_count=1, section_count=2) == "1 title updated"
        assert title_update_detail(title_count=2, section_count=3) == "2 titles updated"

    def test_asr_model_label_uses_basename_for_non_hf_absolute_paths(self) -> None:
        assert asr_model_label("/tmp/models/local-model.bin") == "local-model.bin"

    def test_hf_cache_repo_label_returns_none_for_non_hf_path(self) -> None:
        assert hf_cache_repo_label(Path("/tmp/models/local-model.bin")) is None

    def test_token_usage_detail_returns_blank_without_total_tokens(self) -> None:
        assert token_usage_detail({"input_tokens": 2}) == ""

    def test_resolve_llm_processor_uses_environment_processor_details(self, monkeypatch) -> None:
        reporter = RecordingReporter()
        warnings: list[str] = []
        fake_processor = cast(
            "LLMProcessor", SimpleNamespace(provider_name="openai", model_name="gpt-test")
        )
        monkeypatch.setattr(
            "webinar_transcriber.processor.llm.build_llm_processor_from_env",
            lambda: fake_processor,
        )

        resolved, runtime = resolve_llm_processor(
            enable_llm=True,
            llm_processor=None,
            reporter=reporter,
            warnings=warnings,
            llm_runtime=LLMRuntimeState(),
        )

        assert resolved is fake_processor
        assert runtime.provider_name == "openai"
        assert runtime.model_name == "gpt-test"
        assert runtime.report_status == "disabled"
        assert warnings == []


@pytest.fixture
def llm_success_result(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> tuple[ProcessArtifacts, RecordingReporter]:
    install_basic_windowing(monkeypatch)
    reporter = RecordingReporter()

    class FakeLLMProcessor:
        provider_name = "openai"
        model_name = "test-llm-model"

        def report_polish_plan(self, report) -> LLMReportPolishPlan:
            return LLMReportPolishPlan(section_count=len(report.sections), worker_count=1)

        def polish_report_sections_with_progress(
            self, report, *, progress_callback: Callable[[int], None] | None = None
        ) -> LLMSectionPolishResult:
            assert progress_callback is not None
            progress_callback(len(report.sections))
            return LLMSectionPolishResult(
                section_tldrs={
                    "section-1": "Agenda update and next-step reminder for the draft delivery."
                },
                section_transcripts={
                    "section-1": (
                        "Agenda review, and project status update.\n\n"
                        "Next step: please send the draft by Friday."
                    )
                },
                usage={"input_tokens": 12, "output_tokens": 3, "total_tokens": 15},
                warnings=[EXPECTED_LLM_WARNING],
            )

        def polish_report_metadata(
            self, report, *, section_transcripts: dict[str, str]
        ) -> LLMReportMetadataResult:
            assert "section-1" in section_transcripts
            return LLMReportMetadataResult(
                summary=["Refined summary point."],
                action_items=["Refined action item."],
                section_titles={"section-1": "Refined section title"},
                usage={"input_tokens": 8, "output_tokens": 3, "total_tokens": 11},
            )

    artifacts = process_input(
        FIXTURE_DIR / "sample-audio.mp3",
        output_dir=tmp_path / "llm-run",
        transcriber=TestProcessInput.FakeTranscriber(),
        llm_processor=cast("LLMProcessor", FakeLLMProcessor()),
        enable_llm=True,
        reporter=reporter,
    )

    return artifacts, reporter


class TestProcessInputLlm:
    def test_updates_section_text_from_llm_output(self, llm_success_result) -> None:
        artifacts, _reporter = llm_success_result

        assert artifacts.report.sections[0].tldr == (
            "Agenda update and next-step reminder for the draft delivery."
        )
        assert artifacts.report.sections[0].transcript_text == EXPECTED_LLM_SECTION_TEXT

    def test_updates_summary_and_action_items_from_llm(self, llm_success_result) -> None:
        artifacts, _reporter = llm_success_result

        assert artifacts.report.summary == ["Refined summary point."]
        assert artifacts.report.action_items == ["Refined action item."]
        assert artifacts.report.sections[0].title == "Refined section title"

    def test_aggregates_llm_usage_into_diagnostics(self, llm_success_result) -> None:
        artifacts, reporter = llm_success_result

        assert artifacts.diagnostics.llm_enabled
        assert artifacts.diagnostics.llm_model == "test-llm-model"
        assert artifacts.diagnostics.llm_report_status == "applied"
        assert artifacts.diagnostics.warnings == [EXPECTED_LLM_WARNING]
        assert artifacts.diagnostics.llm_report_usage == EXPECTED_LLM_USAGE
        assert "llm_report_sections" in artifacts.diagnostics.stage_durations_sec
        assert "llm_report_metadata" in artifacts.diagnostics.stage_durations_sec
        assert artifacts.diagnostics.llm_report_latency_sec == pytest.approx(
            artifacts.diagnostics.stage_durations_sec["llm_report_sections"]
            + artifacts.diagnostics.stage_durations_sec["llm_report_metadata"],
            abs=1e-4,
        )
        assert reporter.has_progress_event("start", "llm_report_sections", 1.0, None)
        assert reporter.has_progress_event("advance", "llm_report_sections", 1.0, None)
        assert reporter.has_event("finish", "llm_report_sections", "1 section")
        assert (
            "finish",
            "llm_report",
            "1 summary bullet | 1 action item | 1 TL;DR | 26 tokens",
        ) in reporter.events
        assert (
            "start",
            "llm_report_sections",
            "Polishing section text with LLM (openai | test-llm-model | 1 worker)",
        ) in reporter.events
        assert (
            "start",
            "llm_report",
            "Polishing report summary with LLM (openai | test-llm-model)",
        ) in reporter.events
        assert reporter.warnings == [EXPECTED_LLM_WARNING]

    def test_reports_only_polishable_sections_in_llm_progress(self, tmp_path, monkeypatch) -> None:
        install_basic_windowing(monkeypatch)
        reporter = RecordingReporter()
        monkeypatch.setattr(
            "webinar_transcriber.structure.build_report",
            lambda *_args, **_kwargs: ReportDocument(
                title="Demo",
                source_file="sample-audio.mp3",
                media_type=MediaType.AUDIO,
                detected_language="en",
                sections=[
                    ReportSection(
                        id="section-1",
                        title="Music Interlude",
                        start_sec=0.0,
                        end_sec=1.0,
                        transcript_text=(
                            "Music interlude. The raw transcript is preserved in transcript.json."
                        ),
                        is_interlude=True,
                    ),
                    ReportSection(
                        id="section-2",
                        title="Agenda",
                        start_sec=1.0,
                        end_sec=3.0,
                        transcript_text="Agenda review and project status update.",
                    ),
                ],
            ),
        )

        class FakeLLMProcessor:
            provider_name = "openai"
            model_name = "test-llm-model"

            def report_polish_plan(self, report) -> LLMReportPolishPlan:
                del report
                return LLMReportPolishPlan(section_count=1, worker_count=1, skipped_section_count=1)

            def polish_report_sections_with_progress(
                self, report, *, progress_callback: Callable[[int], None] | None = None
            ) -> LLMSectionPolishResult:
                del report
                assert progress_callback is not None
                progress_callback(1)
                return LLMSectionPolishResult(
                    section_tldrs={"section-2": "Agenda recap."},
                    section_transcripts={
                        "section-1": (
                            "Music interlude. The raw transcript is preserved in transcript.json."
                        ),
                        "section-2": EXPECTED_LLM_SECTION_TEXT,
                    },
                    usage={"input_tokens": 12, "output_tokens": 3, "total_tokens": 15},
                    warnings=[
                        "Skipped LLM section polish for likely music/interlude section section-1."
                    ],
                )

            def polish_report_metadata(
                self, report, *, section_transcripts: dict[str, str]
            ) -> LLMReportMetadataResult:
                assert report.sections[1].id == "section-2"
                assert section_transcripts["section-2"] == EXPECTED_LLM_SECTION_TEXT
                return LLMReportMetadataResult(
                    summary=["Refined summary point."],
                    action_items=["Refined action item."],
                    section_titles={"section-2": "Refined section title"},
                    usage={"input_tokens": 8, "output_tokens": 3, "total_tokens": 11},
                )

        artifacts = process_input(
            FIXTURE_DIR / "sample-audio.mp3",
            output_dir=tmp_path / "llm-interlude-run",
            transcriber=TestProcessInput.FakeTranscriber(),
            llm_processor=cast("LLMProcessor", FakeLLMProcessor()),
            enable_llm=True,
            reporter=reporter,
        )

        assert reporter.has_progress_event("start", "llm_report_sections", 1.0, None)
        assert reporter.progress_stage_events("advance", "llm_report_sections") == [
            ("advance", "llm_report_sections", 1.0, None)
        ]
        assert reporter.has_event(
            "finish", "llm_report_sections", "1 section | 1 skipped interlude"
        )
        assert reporter.has_event(
            "finish", "llm_report", "1 summary bullet | 1 action item | 1 TL;DR | 26 tokens"
        )
        assert artifacts.report.sections[0].transcript_text.startswith("Music interlude.")
        assert artifacts.report.sections[1].title == "Refined section title"

    def test_warns_when_llm_configuration_is_missing(self, tmp_path, monkeypatch) -> None:
        install_basic_windowing(monkeypatch)
        reporter = RecordingReporter()

        with patch(
            "webinar_transcriber.processor.llm.build_llm_processor_from_env",
            side_effect=LLMConfigurationError(
                "Missing required LLM environment variables: OPENAI_API_KEY."
            ),
        ):
            artifacts = process_input(
                FIXTURE_DIR / "sample-audio.mp3",
                output_dir=tmp_path / "llm-fallback-run",
                transcriber=TestProcessInput.FakeTranscriber(),
                enable_llm=True,
                reporter=reporter,
            )

        assert artifacts.diagnostics.llm_enabled
        assert artifacts.diagnostics.llm_model is None
        assert artifacts.diagnostics.llm_report_status == "fallback"
        assert artifacts.diagnostics.warnings == [
            "Missing required LLM environment variables: OPENAI_API_KEY."
        ]
        assert reporter.warnings == ["Missing required LLM environment variables: OPENAI_API_KEY."]

    def test_falls_back_when_report_polish_fails(self, tmp_path, monkeypatch) -> None:
        install_basic_windowing(monkeypatch)
        reporter = RecordingReporter()

        class FakeLLMProcessor:
            provider_name = "openai"
            model_name = "test-llm-model"

            def report_polish_plan(self, report) -> LLMReportPolishPlan:
                return LLMReportPolishPlan(section_count=len(report.sections), worker_count=1)

            def polish_report_sections_with_progress(
                self, report, *, progress_callback: Callable[[int], None] | None = None
            ):
                assert progress_callback is not None
                progress_callback(len(report.sections))
                raise LLMProcessingError("Report polishing failed: backend timeout")

            def polish_report_metadata(self, report, *, section_transcripts: dict[str, str]):
                raise AssertionError("metadata polish should not run after section failure")

        artifacts = process_input(
            FIXTURE_DIR / "sample-audio.mp3",
            output_dir=tmp_path / "llm-report-fallback-run",
            transcriber=TestProcessInput.FakeTranscriber(),
            llm_processor=cast("LLMProcessor", FakeLLMProcessor()),
            enable_llm=True,
            reporter=reporter,
        )

        fallback_finish = ("finish", "llm_report_sections", "openai | test-llm-model | fallback")

        assert artifacts.report.summary != []
        assert artifacts.diagnostics.llm_report_status == "fallback"
        assert "Report polishing failed: backend timeout" in artifacts.diagnostics.warnings
        assert fallback_finish in reporter.events

    def test_records_both_llm_stage_timings_when_metadata_polish_falls_back(
        self, tmp_path, monkeypatch
    ) -> None:
        install_basic_windowing(monkeypatch)
        reporter = RecordingReporter()

        class FakeLLMProcessor:
            provider_name = "openai"
            model_name = "test-llm-model"

            def report_polish_plan(self, report) -> LLMReportPolishPlan:
                return LLMReportPolishPlan(section_count=len(report.sections), worker_count=1)

            def polish_report_sections_with_progress(
                self, report, *, progress_callback: Callable[[int], None] | None = None
            ) -> LLMSectionPolishResult:
                assert progress_callback is not None
                progress_callback(len(report.sections))
                return LLMSectionPolishResult(
                    section_tldrs={},
                    section_transcripts={
                        section.id: section.transcript_text for section in report.sections
                    },
                    usage={"input_tokens": 12, "output_tokens": 3, "total_tokens": 15},
                    warnings=[],
                )

            def polish_report_metadata(
                self, report, *, section_transcripts: dict[str, str]
            ) -> LLMReportMetadataResult:
                del report, section_transcripts
                raise LLMProcessingError("metadata polish failed")

        artifacts = process_input(
            FIXTURE_DIR / "sample-audio.mp3",
            output_dir=tmp_path / "llm-metadata-fallback-run",
            transcriber=TestProcessInput.FakeTranscriber(),
            llm_processor=cast("LLMProcessor", FakeLLMProcessor()),
            enable_llm=True,
            reporter=reporter,
        )

        assert artifacts.diagnostics.llm_report_status == "fallback"
        assert artifacts.diagnostics.warnings == ["metadata polish failed"]
        assert "llm_report_sections" in artifacts.diagnostics.stage_durations_sec
        assert "llm_report_metadata" in artifacts.diagnostics.stage_durations_sec
        assert artifacts.diagnostics.llm_report_latency_sec == pytest.approx(
            artifacts.diagnostics.stage_durations_sec["llm_report_sections"]
            + artifacts.diagnostics.stage_durations_sec["llm_report_metadata"],
            abs=1e-4,
        )
        assert reporter.has_event("finish", "llm_report", "openai | test-llm-model | fallback")
